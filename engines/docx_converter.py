#!/usr/bin/env python3
"""
하이브리드 DOCX to Markdown 변환기
테이블: 헤더만 추출하여 작성 템플릿으로 변환
본문: 서식을 유지하며 마크다운으로 변환
"""

import os
import sys
import re
import logging
import subprocess
import argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import tempfile
import shutil

try:
    import docx
    from docx.table import Table as DocxTable
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("Error: pyyaml not installed. Run: pip install pyyaml")
    sys.exit(1)

try:
    from colorama import init, Fore, Style
    init(autoreset=True)
except ImportError:
    # Fallback if colorama not available
    class Fore:
        RED = GREEN = YELLOW = BLUE = CYAN = MAGENTA = WHITE = RESET = ""
    class Style:
        BRIGHT = DIM = RESET_ALL = ""


@dataclass
class TableHeader:
    """테이블 헤더 정보를 담는 데이터클래스"""
    position: int  # 문서 내 테이블 순서
    headers: List[str]  # 헤더 셀 텍스트 리스트
    column_count: int
    is_guide_box: bool = False  # 작성 요령 등 가이드 박스 여부

    def to_markdown(self, empty_rows: int = 2) -> str:
        """
        마크다운 테이블 템플릿 생성 (헤더 + 빈 행)

        Args:
            empty_rows: 템플릿에 포함할 빈 행 개수

        Returns:
            마크다운 형식의 테이블 문자열
        """
        if not self.headers:
            return ""

        # 헤더 행
        header_row = "| " + " | ".join(self.headers) + " |"

        # 구분선
        separator = "|" + "|".join([" " + "-" * max(3, len(h)) + " " for h in self.headers]) + "|"

        # 빈 데이터 행
        empty_cells = " | ".join([" " * max(3, len(h)) for h in self.headers])
        data_rows = "\n".join(["| " + empty_cells + " |"] * empty_rows)

        return f"{header_row}\n{separator}\n{data_rows}"


class DocxToMarkdownConverter:
    """DOCX를 하이브리드 방식으로 마크다운으로 변환하는 클래스"""

    def __init__(
        self,
        docx_path: str,
        output_path: Optional[str] = None,
        config: Optional[Dict] = None
    ):
        """
        Args:
            docx_path: 변환할 DOCX 파일 경로
            output_path: 출력 마크다운 파일 경로 (기본값: 입력 파일명.md)
            config: 설정 딕셔너리
        """
        self.docx_path = Path(docx_path).resolve()
        self.output_path = Path(output_path).resolve() if output_path else \
                          self.docx_path.with_suffix('.md')
        self.config = config or self._load_default_config()
        self.logger = self._setup_logging()
        self.temp_dir = None

        # 설정값 추출
        conv_config = self.config.get('docx_converter', {})
        self.pandoc_path = os.path.expandvars(conv_config.get('pandoc_path', 'pandoc'))
        self.extract_media = conv_config.get('extract_media', True)
        self.media_dir_name = conv_config.get('media_dir_name', 'media')
        self.table_config = conv_config.get('table_extraction', {})
        self.output_config = conv_config.get('output', {})

        self._validate_inputs()

    def _load_default_config(self) -> Dict:
        """config.yaml 로드 또는 기본 설정 반환"""
        config_path = Path(__file__).parent / "config.yaml"
        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f)
            except Exception as e:
                print(f"Warning: Could not load config.yaml: {e}")

        # 기본 설정
        return {
            'docx_converter': {
                'pandoc_path': 'pandoc',
                'extract_media': True,
                'media_dir_name': 'media',
                'table_extraction': {
                    'only_headers': True,
                    'empty_rows': 2,
                    'column_alignment': 'left'
                },
                'output': {
                    'add_metadata_comment': True,
                    'preserve_formatting': True,
                    'overwrite_existing': False
                }
            }
        }

    def _setup_logging(self) -> logging.Logger:
        """로깅 설정"""
        logger = logging.getLogger('docx_converter')
        logger.setLevel(logging.INFO)

        # 콘솔 핸들러
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)

        # 로그 포맷
        formatter = logging.Formatter(
            f'{Fore.CYAN}[%(levelname)s]{Style.RESET_ALL} %(asctime)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        console_handler.setFormatter(formatter)

        logger.addHandler(console_handler)

        # 파일 핸들러 (선택적)
        log_dir = Path(__file__).parent / "logs"
        if log_dir.exists() or True:  # 항상 로그 디렉토리 생성 시도
            log_dir.mkdir(exist_ok=True)
            file_handler = logging.FileHandler(
                log_dir / "docx_converter.log",
                encoding='utf-8'
            )
            file_handler.setLevel(logging.DEBUG)
            file_formatter = logging.Formatter(
                '[%(levelname)s] %(asctime)s - %(message)s',
                datefmt='%Y-%m-%d %H:%M:%S'
            )
            file_handler.setFormatter(file_formatter)
            logger.addHandler(file_handler)

        return logger

    def _validate_inputs(self):
        """입력 검증"""
        # DOCX 파일 존재 확인
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {self.docx_path}")

        if not self.docx_path.suffix.lower() == '.docx':
            raise ValueError(f"File must be .docx format: {self.docx_path}")

        # DOCX 파일 읽기 가능 여부 확인
        try:
            docx.Document(str(self.docx_path))
        except Exception as e:
            raise ValueError(f"Cannot open DOCX file (corrupted or password-protected?): {e}")

        # 출력 디렉토리 쓰기 가능 확인
        output_dir = self.output_path.parent
        if not output_dir.exists():
            try:
                output_dir.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                raise PermissionError(f"Cannot create output directory: {e}")

        # 출력 파일 덮어쓰기 확인
        if self.output_path.exists() and not self.output_config.get('overwrite_existing', False):
            # 파일명에 _converted 추가
            base = self.output_path.stem
            self.output_path = self.output_path.with_stem(f"{base}_converted")
            self.logger.info(f"Output file exists. Using: {self.output_path.name}")

        # Pandoc 실행 가능 여부 확인
        try:
            subprocess.run(
                [self.pandoc_path, '--version'],
                capture_output=True,
                check=True,
                timeout=5
            )
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
            self.logger.warning(f"Pandoc not accessible at {self.pandoc_path}. Will try fallback.")
            # Pandoc이 없어도 계속 진행 (python-docx로만 변환)

    def convert(self) -> Path:
        """
        메인 변환 함수

        Returns:
            생성된 마크다운 파일 경로
        """
        self.logger.info(f"{Fore.GREEN}Starting conversion:{Style.RESET_ALL} {self.docx_path.name}")

        try:
            # 임시 디렉토리 생성
            self.temp_dir = Path(tempfile.mkdtemp(prefix='docx_conv_'))
            self.logger.debug(f"Temp directory: {self.temp_dir}")

            # Phase 1: 테이블 헤더 추출
            self.logger.info(f"{Fore.YELLOW}Phase 1:{Style.RESET_ALL} Extracting table headers...")
            table_headers = self._extract_table_headers()
            self.logger.info(f"  → Extracted {len(table_headers)} table headers")

            # Phase 2: Pandoc으로 본문 변환
            self.logger.info(f"{Fore.YELLOW}Phase 2:{Style.RESET_ALL} Converting body with Pandoc...")
            pandoc_md = self._convert_body_with_pandoc()
            if pandoc_md:
                self.logger.info(f"  → Pandoc conversion completed ({len(pandoc_md)} characters)")
            else:
                self.logger.warning("  → Pandoc conversion failed, using fallback")
                pandoc_md = self._fallback_conversion()

            # Phase 3: 콘텐츠 병합
            self.logger.info(f"{Fore.YELLOW}Phase 3:{Style.RESET_ALL} Merging content...")
            merged_content = self._merge_content(pandoc_md, table_headers)
            line_count = len(merged_content.splitlines())
            self.logger.info(f"  → Merged content: {line_count} lines")

            # Phase 4: 출력 저장
            self.logger.info(f"{Fore.YELLOW}Phase 4:{Style.RESET_ALL} Saving output...")
            self._save_output(merged_content)
            self.logger.info(f"  → {Fore.GREEN}Output saved:{Style.RESET_ALL} {self.output_path}")

            # Cleanup
            self._cleanup_temp_files()

            self.logger.info(f"{Fore.GREEN}[OK] Conversion completed successfully!{Style.RESET_ALL}")
            return self.output_path

        except Exception as e:
            self.logger.error(f"{Fore.RED}Conversion failed:{Style.RESET_ALL} {e}")
            # 임시 파일 유지 (디버깅용)
            if self.temp_dir:
                self.logger.info(f"Temp files kept for debugging: {self.temp_dir}")
            raise

    def _remove_duplicate_headers(self, headers: List[str]) -> List[str]:
        """
        병합된 셀로 인한 중복 헤더 제거

        Args:
            headers: 원본 헤더 리스트

        Returns:
            중복이 제거된 헤더 리스트
        """
        if not headers:
            return headers

        result = []
        prev_text = None
        consecutive_count = 0

        for header in headers:
            text = header.strip()

            if text and text == prev_text:
                # 같은 텍스트가 연속으로 나타남
                consecutive_count += 1
                # 3번 이상 연속이면 병합 셀로 간주하고 건너뜀
                if consecutive_count >= 3:
                    continue
            else:
                consecutive_count = 1
                prev_text = text

            result.append(header)

        return result

    def _extract_table_headers(self) -> List[TableHeader]:
        """
        DOCX에서 테이블 헤더만 추출

        Returns:
            TableHeader 객체 리스트
        """
        try:
            doc = docx.Document(str(self.docx_path))
        except Exception as e:
            self.logger.error(f"Failed to open DOCX with python-docx: {e}")
            return []

        table_headers = []
        guide_keywords = ['작성 요령', '작성요령', '작성 방법', '작성방법']

        for idx, table in enumerate(doc.tables):
            try:
                if not table.rows:
                    self.logger.warning(f"Table {idx} has no rows, skipping")
                    continue

                # 첫 번째 행을 헤더로 간주
                first_row = table.rows[0]
                headers = []

                for cell in first_row.cells:
                    # 셀 텍스트 추출 (여러 단락이 있을 수 있음)
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    headers.append(cell_text or " ")  # 빈 셀은 공백으로

                if not headers:
                    self.logger.warning(f"Table {idx} has no header cells, skipping")
                    continue

                # 병합 셀 중복 제거
                original_count = len(headers)
                headers = self._remove_duplicate_headers(headers)
                if len(headers) < original_count:
                    self.logger.debug(f"Table {idx}: Removed {original_count - len(headers)} duplicate columns ({original_count} → {len(headers)})")

                # 가이드 박스 감지: 첫 번째 셀에 "작성 요령" 키워드 포함 여부
                first_cell_text = headers[0] if headers else ""
                is_guide = any(keyword in first_cell_text for keyword in guide_keywords)

                table_header = TableHeader(
                    position=idx,
                    headers=headers,
                    column_count=len(headers),
                    is_guide_box=is_guide
                )
                table_headers.append(table_header)

                guide_marker = " [GUIDE BOX]" if is_guide else ""
                self.logger.debug(f"Table {idx}{guide_marker}: {len(headers)} columns - {headers[:3]}...")

            except Exception as e:
                self.logger.warning(f"Failed to extract headers from table {idx}: {e}")
                continue

        return table_headers

    def _convert_body_with_pandoc(self) -> Optional[str]:
        """
        Pandoc으로 DOCX를 마크다운으로 변환

        Returns:
            변환된 마크다운 텍스트 또는 None (실패 시)
        """
        temp_md = self.temp_dir / "pandoc_output.md"
        media_dir = self.output_path.parent / self.media_dir_name

        # Pandoc 명령
        pandoc_cmd = [
            self.pandoc_path,
            str(self.docx_path),
            '-o', str(temp_md),
            '--wrap=none',
            '--markdown-headings=atx'
        ]

        # 미디어 추출 옵션
        if self.extract_media:
            pandoc_cmd.extend(['--extract-media', str(media_dir.parent)])

        try:
            result = subprocess.run(
                pandoc_cmd,
                capture_output=True,
                text=True,
                encoding='utf-8',
                timeout=60,
                check=True
            )

            if result.stderr:
                self.logger.debug(f"Pandoc stderr: {result.stderr}")

            # 출력 파일 읽기
            if temp_md.exists():
                with open(temp_md, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                self.logger.error("Pandoc did not create output file")
                return None

        except subprocess.TimeoutExpired:
            self.logger.error("Pandoc conversion timed out")
            return None
        except subprocess.CalledProcessError as e:
            self.logger.error(f"Pandoc conversion failed: {e.stderr}")
            return None
        except Exception as e:
            self.logger.error(f"Pandoc conversion error: {e}")
            return None

    def _fallback_conversion(self) -> str:
        """
        Pandoc 실패 시 python-docx만으로 기본 변환

        Returns:
            기본 마크다운 텍스트
        """
        self.logger.info("Using fallback conversion (python-docx only)...")

        try:
            doc = docx.Document(str(self.docx_path))
        except Exception as e:
            self.logger.error(f"Fallback conversion failed: {e}")
            return "# Conversion Failed\n\nCould not convert document."

        lines = []

        for element in doc.element.body:
            # 단락 처리
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            # 스타일에 따라 마크다운 형식 적용 (기본)
                            if para.style.name.startswith('Heading'):
                                level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                                lines.append(f"{'#' * int(level)} {text}")
                            else:
                                lines.append(text)
                            lines.append("")
                        break

            # 테이블 처리 (전체 테이블 - 헤더뿐 아니라 모든 내용)
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        # 이 테이블은 나중에 헤더로 교체될 것
                        lines.append(f"[TABLE_{doc.tables.index(table)}]")
                        lines.append("")
                        break

        return "\n".join(lines)

    def _merge_content(self, pandoc_md: str, table_headers: List[TableHeader]) -> str:
        """
        Pandoc 출력과 추출된 테이블 헤더를 병합

        Args:
            pandoc_md: Pandoc이 생성한 마크다운
            table_headers: 추출된 테이블 헤더 리스트

        Returns:
            병합된 최종 마크다운
        """
        if not table_headers:
            # 테이블이 없으면 그대로 반환
            return self._add_metadata(pandoc_md)

        # 테이블 교체
        lines = pandoc_md.split('\n')
        result_lines = []
        table_idx = 0
        in_table = False
        table_buffer = []

        empty_rows = self.table_config.get('empty_rows', 2)

        for line in lines:
            stripped = line.strip()

            # 테이블 시작 감지 (| 또는 + 로 시작하는 행)
            # Pandoc은 pipe table (|)과 grid table (+)을 생성할 수 있음
            is_table_line = (
                stripped.startswith('|') or
                stripped.startswith('+') or
                (stripped.startswith(':') and ':' in stripped and '|' in stripped)  # alignment row
            )

            if is_table_line:
                if not in_table:
                    in_table = True
                    table_buffer = [line]
                else:
                    table_buffer.append(line)
            else:
                # 테이블 종료
                if in_table:
                    in_table = False

                    # 테이블 처리: 가이드 박스는 원본 유지, 일반 테이블은 헤더만 추출
                    if table_idx < len(table_headers):
                        current_table = table_headers[table_idx]

                        if current_table.is_guide_box:
                            # 가이드 박스는 Pandoc 원본 유지
                            result_lines.extend(table_buffer)
                            self.logger.debug(f"Preserved guide box (table {table_idx})")
                        else:
                            # 일반 테이블은 헤더 템플릿으로 교체
                            header_template = current_table.to_markdown(empty_rows)
                            result_lines.append(header_template)

                        result_lines.append("")  # 테이블 후 빈 줄 추가
                        table_idx += 1
                    else:
                        # 추출된 헤더보다 많은 테이블이 있으면 원본 유지
                        result_lines.extend(table_buffer)
                        result_lines.append("")

                    table_buffer = []

                    # 현재 줄이 빈 줄이 아니면 추가
                    if stripped:
                        result_lines.append(line)
                else:
                    result_lines.append(line)

        # 마지막에 테이블이 있었다면 처리
        if in_table and table_buffer:
            if table_idx < len(table_headers):
                current_table = table_headers[table_idx]

                if current_table.is_guide_box:
                    # 가이드 박스는 원본 유지
                    result_lines.extend(table_buffer)
                else:
                    # 일반 테이블은 헤더 템플릿으로 교체
                    header_template = current_table.to_markdown(empty_rows)
                    result_lines.append(header_template)
            else:
                result_lines.extend(table_buffer)

        merged = '\n'.join(result_lines)

        # 메타데이터 추가
        return self._add_metadata(merged)

    def _add_metadata(self, content: str) -> str:
        """
        마크다운 상단에 메타데이터 주석 추가

        Args:
            content: 마크다운 내용

        Returns:
            메타데이터가 추가된 마크다운
        """
        if not self.output_config.get('add_metadata_comment', True):
            return content

        metadata = f"""<!--
Converted from: {self.docx_path.name}
Conversion date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Purpose: Writing template - fill in table cells
Converter: docx_converter.py (Hybrid mode)
-->

"""
        return metadata + content

    def _save_output(self, content: str):
        """
        최종 마크다운을 파일로 저장

        Args:
            content: 저장할 마크다운 내용
        """
        try:
            with open(self.output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            raise IOError(f"Failed to save output file: {e}")

    def _cleanup_temp_files(self):
        """임시 파일 정리"""
        if self.temp_dir and self.temp_dir.exists():
            try:
                shutil.rmtree(self.temp_dir)
                self.logger.debug("Cleaned up temporary files")
            except Exception as e:
                self.logger.warning(f"Could not clean up temp files: {e}")


def setup_cli() -> argparse.ArgumentParser:
    """CLI 인터페이스 설정"""
    parser = argparse.ArgumentParser(
        description='하이브리드 DOCX to Markdown 변환기 (테이블 헤더 템플릿)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예제:
  # 기본 변환 (파일명.md로 저장)
  python docx_converter.py "proposal.docx"

  # 출력 경로 지정
  python docx_converter.py "proposal.docx" -o "output/proposal.md"

  # 디렉토리 내 모든 DOCX 파일 변환
  python docx_converter.py --batch "1. Projects/2026_ACE 과제"
        """
    )

    parser.add_argument(
        'input',
        help='변환할 DOCX 파일 경로 (또는 --batch와 함께 디렉토리)'
    )

    parser.add_argument(
        '-o', '--output',
        help='출력 마크다운 파일 경로 (기본값: 입력 파일명.md)'
    )

    parser.add_argument(
        '--batch',
        action='store_true',
        help='디렉토리 내 모든 DOCX 파일 변환'
    )

    parser.add_argument(
        '--config',
        help='설정 파일 경로 (기본값: scripts/config.yaml)'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='상세 로그 출력'
    )

    return parser


def main():
    """메인 함수"""
    parser = setup_cli()
    args = parser.parse_args()

    # 설정 로드
    config = None
    if args.config:
        config_path = Path(args.config)
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = yaml.safe_load(f)

    # 배치 모드
    if args.batch:
        input_dir = Path(args.input)
        if not input_dir.is_dir():
            print(f"Error: {args.input} is not a directory")
            return 1

        docx_files = list(input_dir.glob('**/*.docx'))
        if not docx_files:
            print(f"No DOCX files found in {input_dir}")
            return 0

        print(f"Found {len(docx_files)} DOCX files to convert")

        success_count = 0
        for docx_file in docx_files:
            try:
                print(f"\n{Fore.CYAN}Converting:{Style.RESET_ALL} {docx_file.name}")
                converter = DocxToMarkdownConverter(
                    docx_path=str(docx_file),
                    config=config
                )
                converter.convert()
                success_count += 1
            except Exception as e:
                print(f"{Fore.RED}Failed:{Style.RESET_ALL} {e}")

        print(f"\n{Fore.GREEN}[Batch Complete]{Style.RESET_ALL} {success_count}/{len(docx_files)} successful")
        return 0

    # 단일 파일 모드
    try:
        converter = DocxToMarkdownConverter(
            docx_path=args.input,
            output_path=args.output,
            config=config
        )

        if args.verbose:
            converter.logger.setLevel(logging.DEBUG)
            for handler in converter.logger.handlers:
                handler.setLevel(logging.DEBUG)

        output_path = converter.convert()
        print(f"\n{Fore.GREEN}[Success]{Style.RESET_ALL} Output: {output_path}")
        return 0

    except Exception as e:
        print(f"\n{Fore.RED}[Error]{Style.RESET_ALL} {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
